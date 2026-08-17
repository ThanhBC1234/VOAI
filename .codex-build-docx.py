from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "docs" / "AI_FIX_ALL_ISSUES.md"
OUTPUT = ROOT / "docs" / "AI_FIX_ALL_ISSUES.docx"

FONT_BODY = "Calibri"
FONT_CODE = "Consolas"
COLOR_BODY = "1F2937"
COLOR_BLUE = "2E74B5"
COLOR_DARK_BLUE = "1F4D78"
COLOR_NAVY = "16324F"
COLOR_MUTED = "667085"
COLOR_LIGHT_BORDER = "D6DEE8"
COLOR_TABLE_HEAD = "E8EEF5"
COLOR_CALLOUT = "F4F6F9"
COLOR_P1 = "9B1C1C"
COLOR_P2 = "7A5A00"
COLOR_P3 = "1F4D78"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = FONT_BODY,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), COLOR_LIGHT_BORDER)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def paragraph_shading(paragraph, fill: str, *, left_border: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    if left_border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        p_bdr.append(left)


def paragraph_bottom_border(paragraph, color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=COLOR_MUTED)


def add_hyperlink(paragraph, url: str, display: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLOR_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_BODY)
    r_pr.extend([r_fonts, color, underline])
    text = OxmlElement("w:t")
    text.text = display
    run.extend([r_pr, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def link_display(url: str) -> str:
    if "nbformat.readthedocs.io" in url:
        return "Đặc tả Jupyter nbformat về cell ID"
    if "github.blog" in url and "node-20" in url:
        return "Thông báo GitHub về Node 20"
    return url


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|<https?://[^>]+>)")


def add_inline(paragraph, text: str, *, default_size: float | None = None):
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=default_size, color=COLOR_BODY)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=FONT_CODE, size=9.2 if default_size is None else min(default_size, 9.2), color=COLOR_DARK_BLUE)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=COLOR_BODY, bold=True)
        else:
            url = token[1:-1]
            add_hyperlink(paragraph, url, link_display(url))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=default_size, color=COLOR_BODY)


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLOR_BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, COLOR_BLUE, 18, 10),
        "Heading 2": (13, COLOR_BLUE, 14, 7),
        "Heading 3": (12, COLOR_DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    list_style = styles["List Paragraph"]
    list_style.font.name = FONT_BODY
    list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    list_style.font.size = Pt(11)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = FONT_CODE
    code._element.rPr.rFonts.set(qn("w:ascii"), FONT_CODE)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CODE)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE)
    code.font.size = Pt(9)
    code.font.color.rgb = rgb("243247")
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.05


def add_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    for level in range(3):
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def add_numbering(doc: Document, *, bullet: bool) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        if bullet:
            lvl_text.set(qn("w:val"), "•" if level == 0 else "◦")
        else:
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        marker = 260 + level * 360
        text_indent = 540 + level * 360
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(text_indent))
        indent.set(qn("w:hanging"), str(text_indent - marker))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, indent, spacing])
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT_BODY)
        r_fonts.set(qn("w:hAnsi"), FONT_BODY)
        r_pr.append(r_fonts)
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
        abstract.append(lvl)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    return abstract_id, add_num_instance(doc, abstract_id)


def apply_numbering(paragraph, num_id: int, level: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(86)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VOAI LAB")
    set_run_font(run, size=11, color=COLOR_BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("BẢN GIAO VIỆC SỬA TOÀN BỘ LỖI")
    set_run_font(run, size=27, color=COLOR_NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("Runbook kỹ thuật dành cho AI/kỹ sư tiếp nhận")
    set_run_font(run, size=13, color=COLOR_MUTED, italic=True)

    metrics = doc.add_table(rows=2, cols=3)
    values = [("7", "TICKET P1"), ("12", "TICKET P2"), ("6", "TICKET P3")]
    for idx, (value, label) in enumerate(values):
        cell = metrics.cell(0, idx)
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=19, color=COLOR_NAVY, bold=True)
        cell = metrics.cell(1, idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=8.5, color=COLOR_MUTED, bold=True)
    set_table_geometry(metrics, [3120, 3120, 3120])

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(22)
    info.paragraph_format.space_after = Pt(16)
    add_inline(
        info,
        "Audit: 15/08/2026  |  Nhánh: main  |  Commit: `cae07fc024ad4fd4b74bc211ec4eddad0196c9ae`",
        default_size=9.5,
    )

    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.22)
    callout.paragraph_format.right_indent = Inches(0.22)
    callout.paragraph_format.space_before = Pt(10)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.2
    paragraph_shading(callout, COLOR_CALLOUT, left_border=COLOR_BLUE)
    run = callout.add_run("MỤC TIÊU  ")
    set_run_font(run, size=9.5, color=COLOR_DARK_BLUE, bold=True)
    run = callout.add_run(
        "Sửa toàn bộ lỗi đã xác nhận, bổ sung test hồi quy và chỉ tuyên bố hoàn tất khi mọi tiêu chí Definition of Done đều đạt."
    )
    set_run_font(run, size=10.2, color=COLOR_BODY)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Tài liệu nguồn sự thật cho phạm vi sửa lỗi")
    set_run_font(run, size=9.5, color=COLOR_MUTED, italic=True)

    doc.add_page_break()


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header_p = section.header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(3)
    run = header_p.add_run("VOAI LAB  |  BẢN GIAO VIỆC SỬA LỖI")
    set_run_font(run, size=8.5, color=COLOR_MUTED, bold=True)
    paragraph_bottom_border(header_p, COLOR_LIGHT_BORDER)

    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_footer.add_run("VOAI Lab  •  Audit 15/08/2026  •  Phiên bản 1.0")
    set_run_font(run, size=8.5, color=COLOR_MUTED)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.25))
    run = footer.add_run("Commit cae07fc  |  Tài liệu bàn giao")
    set_run_font(run, size=8.5, color=COLOR_MUTED)
    run = footer.add_run("\tTrang ")
    set_run_font(run, size=8.5, color=COLOR_MUTED)
    add_page_field(footer)


def add_heading(doc: Document, text: str, level: int, *, page_break_before: bool = False):
    mapped = min(max(level - 1, 1), 3)
    clean = text.strip()
    if clean.startswith("[ ] "):
        clean = "☐ " + clean[4:]
    p = doc.add_paragraph(style=f"Heading {mapped}")
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.keep_with_next = True
    add_inline(p, clean)
    if "-P1-" in clean:
        for run in p.runs:
            run.font.color.rgb = rgb(COLOR_P1)
    elif "-P2-" in clean:
        for run in p.runs:
            run.font.color.rgb = rgb(COLOR_P2)
    elif "-P3-" in clean:
        for run in p.runs:
            run.font.color.rgb = rgb(COLOR_P3)
    return p


def add_callout(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    paragraph_shading(p, COLOR_CALLOUT, left_border=COLOR_BLUE)
    add_inline(p, text, default_size=10.3)


def add_code_block(doc: Document, code_lines: list[str]):
    p = doc.add_paragraph(style="Code Block")
    paragraph_shading(p, "F2F4F7", left_border="98A2B3")
    for idx, line in enumerate(code_lines):
        run = p.add_run(line)
        set_run_font(run, name=FONT_CODE, size=9, color="243247")
        if idx < len(code_lines) - 1:
            run.add_break()


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    if cols == 3:
        widths = [560, 3150, 5650]
    elif cols == 2:
        widths = [2700, 6660]
    else:
        base = CONTENT_WIDTH_DXA // cols
        widths = [base] * cols
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)

    for row_idx, values in enumerate(rows):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, values[col_idx] if col_idx < len(values) else "", default_size=9.1)
            if row_idx == 0:
                set_cell_shading(cell, COLOR_TABLE_HEAD)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(COLOR_DARK_BLUE)
    mark_header_row(table.rows[0])
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def is_special_start(line: str) -> bool:
    stripped = line.lstrip()
    return (
        not stripped
        or stripped.startswith("#")
        or stripped.startswith(">")
        or stripped.startswith("```")
        or stripped == "---"
        or stripped.startswith("|")
        or re.match(r"^(?:[-*]\s+|\d+\.\s+)", stripped) is not None
    )


def parse_markdown(
    doc: Document,
    text: str,
    bullet_num_id: int,
    decimal_abstract_id: int,
):
    lines = text.splitlines()
    index = 0
    skipped_title = False
    skipped_initial_quote = False
    last_list_kind: str | None = None
    active_decimal_num_id: int | None = None

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            last_list_kind = None
            index += 1
            continue

        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue

        if stripped.startswith(">") and not skipped_initial_quote:
            quote_parts = []
            while index < len(lines) and lines[index].lstrip().startswith(">"):
                quote_parts.append(lines[index].lstrip()[1:].strip())
                index += 1
            skipped_initial_quote = True
            continue

        if stripped == "---":
            last_list_kind = None
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            add_code_block(doc, code_lines)
            last_list_kind = None
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            add_heading(
                doc,
                heading.group(2),
                len(heading.group(1)),
                page_break_before=False,
            )
            last_list_kind = None
            index += 1
            continue

        if stripped.startswith(">"):
            parts = []
            while index < len(lines) and lines[index].lstrip().startswith(">"):
                parts.append(lines[index].lstrip()[1:].strip())
                index += 1
            add_callout(doc, " ".join(parts))
            last_list_kind = None
            continue

        if stripped.startswith("|"):
            raw_rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                raw = lines[index].strip().strip("|")
                cells = [cell.strip() for cell in raw.split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
                    raw_rows.append(cells)
                index += 1
            add_markdown_table(doc, raw_rows)
            last_list_kind = None
            continue

        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            indent = len(list_match.group(1).replace("\t", "    "))
            level = min(indent // 2, 2)
            marker = list_match.group(2)
            item_text = list_match.group(3).strip()
            index += 1
            continuation = []
            while index < len(lines):
                nxt = lines[index]
                if not nxt.strip():
                    break
                if is_special_start(nxt):
                    break
                if len(nxt) - len(nxt.lstrip()) > indent:
                    continuation.append(nxt.strip())
                    index += 1
                else:
                    break
            if continuation:
                item_text += " " + " ".join(continuation)
            is_checkbox = item_text.startswith("[ ] ")
            if is_checkbox:
                item_text = "☐ " + item_text[4:]
            p = doc.add_paragraph(style="List Paragraph")
            if is_checkbox:
                p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
                p.paragraph_format.first_line_indent = Inches(-0.188)
                last_list_kind = "checkbox"
            elif marker in ("-", "*"):
                apply_numbering(p, bullet_num_id, level)
                last_list_kind = "bullet"
            else:
                if last_list_kind != "decimal" or active_decimal_num_id is None:
                    active_decimal_num_id = add_num_instance(doc, decimal_abstract_id)
                apply_numbering(p, active_decimal_num_id, level)
                last_list_kind = "decimal"
            add_inline(p, item_text)
            continue

        parts = [stripped]
        index += 1
        while index < len(lines) and lines[index].strip() and not is_special_start(lines[index]):
            parts.append(lines[index].strip())
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(parts))
        last_list_kind = None


def set_update_fields_on_open(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    _, bullet_num_id = add_numbering(doc, bullet=True)
    decimal_abstract_id, _ = add_numbering(doc, bullet=False)
    set_update_fields_on_open(doc)

    props = doc.core_properties
    props.title = "Bản giao việc sửa toàn bộ lỗi VOAI Lab"
    props.subject = "Runbook kỹ thuật cho AI/kỹ sư tiếp nhận"
    props.author = "VOAI Lab"
    props.keywords = "VOAI, audit, bugfix, runbook, acceptance tests"
    props.comments = "Sinh từ docs/AI_FIX_ALL_ISSUES.md tại commit cae07fc."

    add_cover(doc)
    parse_markdown(
        doc,
        SOURCE.read_text(encoding="utf-8"),
        bullet_num_id,
        decimal_abstract_id,
    )
    doc.save(OUTPUT)
    print(f"Created {OUTPUT} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
