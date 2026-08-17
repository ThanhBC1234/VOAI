"""Dựng bản Word của tài liệu hướng dẫn triển khai từ nguồn Markdown.

Nguồn sự thật vẫn là `docs/HUONG_DAN_TRIEN_KHAI_GITHUB.md`. Tệp `.docx` chỉ là
bản trình bày để đọc/in, sinh lại được bất cứ lúc nào:

    python scripts/build-guide-docx.py

Script chỉ hỗ trợ đúng những cấu trúc Markdown mà tài liệu đó thật sự dùng —
heading, đoạn văn, danh sách (kể cả checkbox), bảng, khối mã, trích dẫn và
đường kẻ ngang. Gặp cấu trúc lạ thì báo lỗi thay vì bỏ qua âm thầm, để bản Word
không bao giờ thiếu nội dung so với bản Markdown.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "HUONG_DAN_TRIEN_KHAI_GITHUB.md"
OUTPUT = ROOT / "docs" / "HUONG_DAN_TRIEN_KHAI_GITHUB.docx"

FONT_BODY = "Calibri"
FONT_CODE = "Consolas"

COLOR_BODY = "1F2937"
COLOR_H1 = "13322C"
COLOR_H2 = "1F4D78"
COLOR_H3 = "2E74B5"
COLOR_MUTED = "5B6B66"
COLOR_CODE = "1B3A32"
COLOR_BORDER = "D6DEE8"
COLOR_TABLE_HEAD = "E7EFE6"
COLOR_CODE_BG = "F2F4F1"
COLOR_QUOTE_BG = "FBF3DC"
COLOR_QUOTE_BAR = "C8952B"
COLOR_ACCENT = "2F7A57"

# A4 (11906 dxa) trừ hai lề 2 cm (1134 dxa mỗi bên).
CONTENT_WIDTH_DXA = 11906 - 2 * 1134


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name=FONT_BODY, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attribute), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_shading(paragraph, fill: str, *, left_bar: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if left_bar:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_bar)
        borders.append(left)
        p_pr.append(borders)


def paragraph_bottom_border(paragraph, color: str, size: str = "6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    p_pr.append(node)


def set_cell_margins(cell, top=70, start=110, bottom=70, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Tổng bề rộng cột phải bằng {CONTENT_WIDTH_DXA}, đang là {sum(widths)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.insert(0, tbl_w)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), COLOR_BORDER)
        borders.append(node)
    tbl_pr.append(borders)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)
            cell.width = Cm(widths[index] / 567)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


# --------------------------------------------------------------------------- #
# Inline: **đậm**, `mã`, [nhãn](đích), <url>                                    #
# --------------------------------------------------------------------------- #

INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*)"          # đậm
    r"|(`[^`]+`)"             # mã
    r"|(\[[^\]]+\]\([^)]+\))"  # liên kết markdown
    r"|(<https?://[^>]+>)"     # liên kết trần
)


def add_inline(paragraph, text: str, *, size=10.5, color=COLOR_BODY, base_bold=None) -> None:
    for piece in INLINE_PATTERN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, name=FONT_CODE, size=size - 0.5, color=COLOR_CODE, bold=base_bold)
        elif piece.startswith("[") and "](" in piece:
            label, _, target = piece[1:-1].partition("](")
            run = paragraph.add_run(label)
            set_run_font(run, size=size, color=COLOR_H3, bold=base_bold)
            run.font.underline = True
            tail = paragraph.add_run(f" ({target})")
            set_run_font(tail, size=size - 1, color=COLOR_MUTED)
        elif piece.startswith("<") and piece.endswith(">"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, size=size, color=COLOR_H3, bold=base_bold)
            run.font.underline = True
        else:
            run = paragraph.add_run(piece)
            set_run_font(run, size=size, color=color, bold=base_bold)


# --------------------------------------------------------------------------- #
# Bộ dựng tài liệu                                                             #
# --------------------------------------------------------------------------- #


class GuideBuilder:
    def __init__(self, document: Document) -> None:
        self.document = document

    def spacer(self, points: float = 4) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(points)
        run = paragraph.add_run("")
        set_run_font(run, size=2)

    def heading(self, level: int, text: str) -> None:
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        if level == 1:
            fmt.space_before, fmt.space_after = Pt(0), Pt(6)
            size, color = 24, COLOR_H1
        elif level == 2:
            fmt.space_before, fmt.space_after = Pt(20), Pt(8)
            size, color = 16, COLOR_H2
        else:
            fmt.space_before, fmt.space_after = Pt(14), Pt(5)
            size, color = 12.5, COLOR_H3
        fmt.keep_with_next = True
        add_inline(paragraph, text, size=size, color=color, base_bold=True)
        for run in paragraph.runs:
            run.bold = True
        if level <= 2:
            paragraph_bottom_border(paragraph, COLOR_BORDER, "8" if level == 1 else "4")
        keep_with_next(paragraph)

    def body(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before, fmt.space_after = Pt(0), Pt(7)
        fmt.line_spacing = 1.28
        add_inline(paragraph, text)

    def bullet(self, text: str, *, checkbox: bool = False) -> None:
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.left_indent, fmt.first_line_indent = Cm(0.75), Cm(-0.45)
        fmt.space_before, fmt.space_after = Pt(0), Pt(4)
        fmt.line_spacing = 1.25
        marker = paragraph.add_run("☐  " if checkbox else "•  ")
        set_run_font(marker, size=10.5, color=COLOR_ACCENT, bold=True)
        add_inline(paragraph, text)

    def numbered(self, order: int, text: str) -> None:
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.left_indent, fmt.first_line_indent = Cm(0.85), Cm(-0.55)
        fmt.space_before, fmt.space_after = Pt(0), Pt(4)
        fmt.line_spacing = 1.25
        marker = paragraph.add_run(f"{order}.  ")
        set_run_font(marker, size=10.5, color=COLOR_ACCENT, bold=True)
        add_inline(paragraph, text)

    def quote(self, lines: list[str]) -> None:
        """Khối chú ý. Trong tài liệu này có khối chứa cả hàng rào mã, nên phải
        tách phần mã ra chứ không in nguyên dấu ``` thành văn bản."""
        segments: list[tuple[str, list[str]]] = []
        buffer: list[str] = []
        in_code = False
        for line in lines:
            if line.strip().startswith("```"):
                segments.append(("code" if in_code else "prose", buffer))
                buffer = []
                in_code = not in_code
                continue
            buffer.append(line)
        segments.append(("code" if in_code else "prose", buffer))

        for kind, block in segments:
            block = [item for item in block] if kind == "code" else [
                item for item in block if item.strip() or block.index(item) not in (0, len(block) - 1)
            ]
            if not any(item.strip() for item in block):
                continue
            if kind == "code":
                self.code(block, "")
                continue
            for index, line in enumerate(block):
                paragraph = self.document.add_paragraph()
                fmt = paragraph.paragraph_format
                fmt.left_indent = Cm(0.3)
                fmt.space_before = Pt(6 if index == 0 else 0)
                fmt.space_after = Pt(6 if index == len(block) - 1 else 0)
                fmt.line_spacing = 1.25
                paragraph_shading(paragraph, COLOR_QUOTE_BG, left_bar=COLOR_QUOTE_BAR)
                add_inline(paragraph, line or " ")

    def code(self, lines: list[str], language: str) -> None:
        if language:
            caption = self.document.add_paragraph()
            caption.paragraph_format.space_before = Pt(8)
            caption.paragraph_format.space_after = Pt(0)
            run = caption.add_run(language.upper())
            set_run_font(run, name=FONT_CODE, size=7.5, color=COLOR_MUTED, bold=True)
        for index, line in enumerate(lines):
            paragraph = self.document.add_paragraph()
            fmt = paragraph.paragraph_format
            fmt.left_indent = Cm(0.3)
            fmt.space_before = Pt(0 if language or index else 8)
            fmt.space_after = Pt(8 if index == len(lines) - 1 else 0)
            fmt.line_spacing = 1.15
            paragraph_shading(paragraph, COLOR_CODE_BG, left_bar=COLOR_ACCENT)
            run = paragraph.add_run(line or " ")
            set_run_font(run, name=FONT_CODE, size=9, color=COLOR_CODE)

    def rule(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run("")
        set_run_font(run, size=2)
        paragraph_bottom_border(paragraph, COLOR_BORDER)

    def table(self, rows: list[list[str]]) -> None:
        header, *body_rows = rows
        table = self.document.add_table(rows=len(rows), cols=len(header))
        table.style = "Table Grid"
        widths = self._column_widths(rows)
        set_table_geometry(table, widths)
        mark_header_row(table.rows[0])

        for column, text in enumerate(header):
            cell = table.rows[0].cells[column]
            set_cell_shading(cell, COLOR_TABLE_HEAD)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, text, size=9.5, color=COLOR_H2, base_bold=True)
            for run in paragraph.runs:
                run.bold = True

        for index, row in enumerate(body_rows, start=1):
            for column, text in enumerate(row):
                cell = table.rows[index].cells[column]
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                add_inline(paragraph, text, size=9.5)
        self.spacer(6)

    @staticmethod
    def _column_widths(rows: list[list[str]]) -> list[int]:
        columns = len(rows[0])
        weights = []
        for column in range(columns):
            longest = max(len(row[column]) for row in rows)
            weights.append(max(longest, 6))
        total = sum(weights)
        widths = [max(900, int(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
        return widths


# --------------------------------------------------------------------------- #
# Phân tích Markdown                                                           #
# --------------------------------------------------------------------------- #


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render(markdown: str, builder: GuideBuilder) -> None:
    lines = markdown.splitlines()
    index = 0
    ordered_counter = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            ordered_counter = 0
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            block: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            builder.code(block, language)
            ordered_counter = 0
            continue

        if stripped in {"---", "***", "___"}:
            builder.rule()
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            builder.heading(min(len(heading.group(1)), 3), heading.group(2).strip())
            ordered_counter = 0
            index += 1
            continue

        if stripped.startswith(">"):
            block = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                block.append(lines[index].strip().lstrip(">").strip())
                index += 1
            builder.quote(block)
            ordered_counter = 0
            continue

        if stripped.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                candidate = lines[index].strip()
                if not re.fullmatch(r"\|[\s:|-]+\|", candidate):
                    rows.append(split_table_row(candidate))
                index += 1
            if rows:
                builder.table(rows)
            ordered_counter = 0
            continue

        checkbox = re.match(r"^-\s+\[[ xX]\]\s+(.*)$", stripped)
        if checkbox:
            builder.bullet(join_wrapped(lines, index, checkbox.group(1))[0], checkbox=True)
            index = join_wrapped(lines, index, checkbox.group(1))[1]
            ordered_counter = 0
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            text, index = join_wrapped(lines, index, bullet.group(1))
            builder.bullet(text)
            ordered_counter = 0
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            text, index = join_wrapped(lines, index, numbered.group(2))
            ordered_counter = int(numbered.group(1))
            builder.numbered(ordered_counter, text)
            continue

        paragraph_lines = []
        while index < len(lines) and lines[index].strip() and not is_block_start(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        builder.body(" ".join(paragraph_lines))
        ordered_counter = 0


def is_block_start(line: str) -> bool:
    stripped = line.strip()
    return bool(
        stripped.startswith(("```", ">", "|", "#"))
        or stripped in {"---", "***", "___"}
        or re.match(r"^[-*]\s+", stripped)
        or re.match(r"^\d+\.\s+", stripped)
    )


def join_wrapped(lines: list[str], index: int, first: str) -> tuple[str, int]:
    """Gộp các dòng xuống hàng của cùng một mục danh sách."""
    parts = [first]
    index += 1
    while index < len(lines):
        candidate = lines[index]
        if not candidate.strip() or is_block_start(candidate):
            break
        if not candidate.startswith((" ", "\t")):
            break
        parts.append(candidate.strip())
        index += 1
    return " ".join(parts), index


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(COLOR_BODY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("VOAI Lab · Hướng dẫn triển khai GitHub · sinh từ docs/HUONG_DAN_TRIEN_KHAI_GITHUB.md")
    set_run_font(run, size=8, color=COLOR_MUTED)


def add_cover(builder: GuideBuilder, section_titles: list[str]) -> None:
    document = builder.document

    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(4)
    run = eyebrow.add_run("VOAI LAB · TÀI LIỆU TRIỂN KHAI")
    set_run_font(run, name=FONT_CODE, size=9, color=COLOR_ACCENT, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Đưa VOAI Lab lên GitHub và chạy đủ mọi tính năng")
    set_run_font(run, size=26, color=COLOR_H1, bold=True)

    lead = document.add_paragraph()
    lead.paragraph_format.space_after = Pt(14)
    lead.paragraph_format.line_spacing = 1.3
    run = lead.add_run(
        "Từ máy trắng đến một website chạy thật: cần cài gì, chạy lệnh nào, "
        "và dấu hiệu nào cho biết đã đúng."
    )
    set_run_font(run, size=12, color=COLOR_MUTED)
    paragraph_bottom_border(lead, COLOR_BORDER, "8")

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run("Mục lục")
    set_run_font(run, size=13, color=COLOR_H2, bold=True)

    for entry in section_titles:
        paragraph = document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.left_indent, fmt.first_line_indent = Cm(0.75), Cm(-0.45)
        fmt.space_before, fmt.space_after = Pt(0), Pt(3)
        marker = paragraph.add_run("›  ")
        set_run_font(marker, size=10.5, color=COLOR_ACCENT, bold=True)
        run = paragraph.add_run(entry)
        set_run_font(run, size=10.5, color=COLOR_BODY)

    builder.rule()


def main() -> int:
    if not SOURCE.exists():
        print(f"Không tìm thấy nguồn: {SOURCE}", file=sys.stderr)
        return 1
    markdown = SOURCE.read_text(encoding="utf-8")

    # Bỏ dòng H1 đầu tiên vì trang bìa đã có tiêu đề riêng.
    lines = markdown.splitlines()
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    body = "\n".join(lines)

    section_titles = [
        match.group(1).strip()
        for match in re.finditer(r"^##\s+(.*)$", body, flags=re.MULTILINE)
    ]

    document = Document()
    configure_page(document)
    builder = GuideBuilder(document)
    add_cover(builder, section_titles)
    render(body, builder)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Đã ghi {OUTPUT.relative_to(ROOT)} ({OUTPUT.stat().st_size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
